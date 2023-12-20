import docx

doc = docx.Document("Python開發工具.docx")
print("段落數: ", len(doc.paragraphs))
para = doc.paragraphs[0]
print("連續文字數: ", len(para.runs))
for run in para.runs:
    print(run.text)
para = doc.paragraphs[2]
print("連續文字數: ", len(para.runs))
for i in range(0, len(para.runs)):
    print(i, ":", para.runs[i].text)



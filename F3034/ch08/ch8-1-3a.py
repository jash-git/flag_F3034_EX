import docx

doc = docx.Document("Word文件範例.docx")
doc.add_heading("Python程式設計", level=1)
doc.add_heading("自動化Excel應用", level=2)
doc.save("Word文件範例2.docx")


import docx

doc = docx.Document("Python開發工具.docx")
print("段落數: ", len(doc.paragraphs))
count = 1
for para in doc.paragraphs:
    print(count, ":", para.text)
    count = count + 1



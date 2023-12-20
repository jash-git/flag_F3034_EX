import docx

doc = docx.Document("Python開發工具2.docx")
print("段落數: ", len(doc.paragraphs))
para2 = doc.paragraphs[1]
para3 = doc.paragraphs[2]
run = para2.add_run("-第一版")
run.bold = True
para3.runs[0].italic = True
doc.save("Python開發工具3.docx") 

import docx

doc = docx.Document("Word文件範例2.docx")
para = doc.paragraphs[1]
para.insert_paragraph_before("Python是一種直譯語言。")
para1 = doc.add_paragraph("Python可以使用openpyxl套件來自動化" + 
                          "處理Excel的編輯、讀取、建立、儲存" +
                          "、合併儲存格等相關編輯操作。")
para1.add_run("Spyder")
para1.add_run(", ")
para1.add_run("Anaconda")
doc.save("Word文件範例3.docx")
 

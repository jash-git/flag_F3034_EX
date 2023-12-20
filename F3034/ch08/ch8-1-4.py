import docx

doc = docx.Document("Python開發工具.docx")
print("段落數: ", len(doc.paragraphs))
para = doc.paragraphs[1]
para2 = doc.paragraphs[3]
para.insert_paragraph_before("Python是一種直譯語言。", 
                             style="Heading 3")
doc.add_paragraph("Python程式是使用直譯器執行程式碼。",
                  style="Subtitle")
para2.style = "Quote"
doc.save("Python開發工具2.docx") 



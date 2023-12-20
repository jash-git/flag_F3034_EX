import docx

doc = docx.Document("Python開發工具.docx")
print("段落數: ", len(doc.paragraphs))
for para in doc.paragraphs:
    if para.style.name.startswith('Heading'):
        print(para.text)


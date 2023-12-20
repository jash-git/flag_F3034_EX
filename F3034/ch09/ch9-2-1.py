import docx
from openpyxl import Workbook

doc = docx.Document("Word表格範例.docx")
print("表格數:", len(doc.tables))
wb = Workbook()
for i in range(len(doc.tables)):
    tb = doc.tables[i]
    if i == 0:
        ws = wb.active
        ws.title = "表格"+str(i)
    else:
        ws = wb.create_sheet(title="表格"+str(i), index=i)
    for i in range(len(tb.rows)):
        row = []
        for j in range(len(tb.columns)):
            row.append(tb.cell(i, j).text)
        ws.append(row)
    
wb.save("Word表格範例.xlsx")
wb.close()